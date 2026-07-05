from docx import Document as DocxDocument

from rag_hybrid_search.ingestion.loaders.docx_loader import DocxLoader


def test_load_joins_nonempty_paragraphs(tmp_path):
    path = tmp_path / "notes.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("")
    doc.add_paragraph("Second paragraph")
    doc.save(str(path))

    loaded = DocxLoader().load(str(path))

    assert loaded.format == "docx"
    assert "First paragraph" in loaded.content
    assert "Second paragraph" in loaded.content
